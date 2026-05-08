import os
from datetime import datetime, date
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL

os.makedirs("documentos", exist_ok=True)

AZUL_TECNM = RGBColor(0x1B, 0x39, 0x6A)

class GeneradorDocumentos:

    def generar(self, tipo, alumno, numero):
        if tipo == 'reporte':
            return self.reporte_bimestral(alumno, numero)
        elif tipo == 'eval_actividades':
            return self.evaluacion_actividades(alumno)
        elif tipo == 'eval_cualitativa':
            return self.evaluacion_cualitativa(alumno)
        elif tipo == 'autoeval':
            return self.autoevaluacion(alumno)
        elif tipo == 'liberacion':
            return self.carta_liberacion(alumno)
        return None

    def _encabezado(self, doc):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("TECNOLÓGICO NACIONAL DE MÉXICO")
        run.bold = True
        run.font.size = Pt(13)
        run.font.color.rgb = AZUL_TECNM

        p2 = doc.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r2 = p2.add_run("DEPARTAMENTO DE GESTIÓN TECNOLÓGICA Y VINCULACIÓN")
        r2.bold = True
        r2.font.size = Pt(11)
        r2.font.color.rgb = AZUL_TECNM

        p3 = doc.add_paragraph()
        p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r3 = p3.add_run("OFICINA DE SERVICIO SOCIAL")
        r3.bold = True
        r3.font.size = Pt(11)
        r3.font.color.rgb = AZUL_TECNM

    def _campo(self, doc, label, valor):
        p = doc.add_paragraph()
        r1 = p.add_run(f"{label}: ")
        r1.bold = True
        r1.font.size = Pt(11)
        r2 = p.add_run(str(valor))
        r2.font.size = Pt(11)
        r2.underline = True

    def reporte_bimestral(self, alumno, numero):
        doc = Document()
        self._encabezado(doc)

        titulo = doc.add_paragraph()
        titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        rt = titulo.add_run(f"Reporte Bimestral de Servicio Social — Reporte No. {numero}")
        rt.bold = True
        rt.font.size = Pt(12)
        rt.font.color.rgb = AZUL_TECNM

        doc.add_paragraph()
        nombre_completo = f"{alumno['apellidos']} {alumno['nombre']}"
        self._campo(doc, "Nombre", nombre_completo)
        self._campo(doc, "Carrera", alumno.get('carrera', ''))
        self._campo(doc, "No. de Control", alumno.get('numero_control', ''))

        horas_reporte = 160
        horas_acumuladas = numero * 160
        fecha_inicio = alumno.get('fecha_inicio', date.today())

        self._campo(doc, "Periodo Reportado", f"Bimestre {numero}")
        self._campo(doc, "Dependencia", alumno.get('departamento', ''))
        self._campo(doc, "Programa", alumno.get('programa', ''))
        self._campo(doc, "Total de horas de este reporte", str(horas_reporte))
        self._campo(doc, "Total de horas acumuladas", str(horas_acumuladas))

        doc.add_paragraph()
        p_act = doc.add_paragraph()
        r_act = p_act.add_run("Resumen de actividades: ")
        r_act.bold = True
        doc.add_paragraph("_" * 80)
        doc.add_paragraph("_" * 80)

        doc.add_paragraph()
        tabla = doc.add_table(rows=2, cols=3)
        tabla.style = 'Table Grid'
        celdas = [
            ("Jefe del Departamento", "", "Firma del Interesado"),
            ("", "", "Vo. Bo. Oficina Servicio Social"),
        ]
        for i, fila in enumerate(celdas):
            for j, texto in enumerate(fila):
                celda = tabla.rows[i].cells[j]
                celda.text = texto
                celda.paragraphs[0].runs[0].bold = True if i == 0 else False

        path = f"documentos/Reporte_{alumno['numero_control']}_Bimestre{numero}.docx"
        doc.save(path)
        return path

    def evaluacion_actividades(self, alumno):
        doc = Document()
        self._encabezado(doc)

        titulo = doc.add_paragraph()
        titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        rt = titulo.add_run("Formato de Evaluación de las Actividades")
        rt.bold = True
        rt.font.size = Pt(12)
        rt.font.color.rgb = AZUL_TECNM

        doc.add_paragraph()
        self._campo(doc, "Nombre del prestador", f"{alumno['nombre']} {alumno['apellidos']}")
        self._campo(doc, "Programa", alumno.get('programa', ''))
        self._campo(doc, "Período de realización", f"Inicio: {alumno.get('fecha_inicio', '')}")

        doc.add_paragraph()
        criterios = [
            "Cumple en tiempo y forma con las actividades encomendadas alcanzando los objetivos",
            "Trabaja en equipo y se adapta a nuevas situaciones",
            "Muestra liderazgo en las actividades encomendadas",
            "Organiza su tiempo y trabaja de manera proactiva",
            "Interpreta la realidad y se sensibiliza aportando soluciones a la problemática",
            "Realiza sugerencias innovadoras para beneficio o mejora del programa",
            "Tiene iniciativa para ayudar en las actividades encomendadas y muestra espíritu de servicio",
        ]

        tabla = doc.add_table(rows=len(criterios)+1, cols=7)
        tabla.style = 'Table Grid'
        encabezados = ["No.", "Criterios a evaluar", "Insuficiente", "Suficiente", "Bueno", "Notable", "Excelente"]
        for j, enc in enumerate(encabezados):
            cell = tabla.rows[0].cells[j]
            cell.text = enc
            cell.paragraphs[0].runs[0].bold = True

        for i, criterio in enumerate(criterios):
            fila = tabla.rows[i+1]
            fila.cells[0].text = str(i+1)
            fila.cells[1].text = criterio

        doc.add_paragraph()
        self._campo(doc, "Observaciones", "")
        doc.add_paragraph("_" * 80)

        path = f"documentos/EvalActividades_{alumno['numero_control']}.docx"
        doc.save(path)
        return path

    def evaluacion_cualitativa(self, alumno):
        doc = Document()
        self._encabezado(doc)

        titulo = doc.add_paragraph()
        titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        rt = titulo.add_run("Formato de Evaluación Cualitativa")
        rt.bold = True
        rt.font.size = Pt(12)
        rt.font.color.rgb = AZUL_TECNM

        doc.add_paragraph()
        self._campo(doc, "Nombre del prestador", f"{alumno['nombre']} {alumno['apellidos']}")
        self._campo(doc, "Programa", alumno.get('programa', ''))
        self._campo(doc, "Período de realización", f"Inicio: {alumno.get('fecha_inicio', '')}")

        criterios = [
            "Cumplí en tiempo y forma con las actividades encomendadas alcanzando los objetivos",
            "Trabajé en equipo y me adapté a nuevas situaciones",
            "Mostré liderazgo en las actividades encomendadas",
            "Organicé mi tiempo y trabajé de manera proactiva",
            "Interpreté la realidad y me sensibilicé aportando soluciones a la problemática",
            "Realicé sugerencias innovadoras para beneficio o mejora del programa",
            "Tuve iniciativa para ayudar en las actividades encomendadas y mostré espíritu de servicio",
        ]

        doc.add_paragraph()
        tabla = doc.add_table(rows=len(criterios)+1, cols=7)
        tabla.style = 'Table Grid'
        encabezados = ["No.", "Criterios a evaluar", "Insuficiente", "Suficiente", "Bueno", "Notable", "Excelente"]
        for j, enc in enumerate(encabezados):
            cell = tabla.rows[0].cells[j]
            cell.text = enc
            cell.paragraphs[0].runs[0].bold = True

        for i, criterio in enumerate(criterios):
            tabla.rows[i+1].cells[0].text = str(i+1)
            tabla.rows[i+1].cells[1].text = criterio

        doc.add_paragraph()
        self._campo(doc, "Observaciones", "")
        doc.add_paragraph("_" * 80)

        path = f"documentos/EvalCualitativa_{alumno['numero_control']}.docx"
        doc.save(path)
        return path

    def autoevaluacion(self, alumno):
        doc = Document()
        self._encabezado(doc)

        titulo = doc.add_paragraph()
        titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        rt = titulo.add_run("Formato de Autoevaluación Cualitativa")
        rt.bold = True
        rt.font.size = Pt(12)
        rt.font.color.rgb = AZUL_TECNM

        doc.add_paragraph()
        self._campo(doc, "Nombre del prestador", f"{alumno['nombre']} {alumno['apellidos']}")
        self._campo(doc, "Programa", alumno.get('programa', ''))

        criterios = [
            "¿Consideras importante la realización del Servicio Social?",
            "¿Consideras que las actividades que realizaste son pertinentes a los fines del Servicio Social?",
            "¿Consideras que las actividades que realizaste contribuyen a tu formación integral?",
            "¿Contribuiste en actividades de beneficio social comunitario?",
            "¿Contribuiste en actividades de protección al medio ambiente?",
            "¿Cómo consideras que las competencias adquiridas en la escuela contribuyeron a las actividades?",
            "¿Consideras que sería factible continuar con este proyecto a Residencias Profesionales?",
            "¿Recomendarías a otro estudiante realizar su Servicio Social en esta dependencia?",
        ]

        doc.add_paragraph()
        tabla = doc.add_table(rows=len(criterios)+1, cols=7)
        tabla.style = 'Table Grid'
        encabezados = ["No.", "Criterios a evaluar", "Insuficiente", "Suficiente", "Bueno", "Notable", "Excelente"]
        for j, enc in enumerate(encabezados):
            cell = tabla.rows[0].cells[j]
            cell.text = enc
            cell.paragraphs[0].runs[0].bold = True

        for i, criterio in enumerate(criterios):
            tabla.rows[i+1].cells[0].text = str(i+1)
            tabla.rows[i+1].cells[1].text = criterio

        doc.add_paragraph()
        self._campo(doc, "Observaciones", "")

        path = f"documentos/Autoevaluacion_{alumno['numero_control']}.docx"
        doc.save(path)
        return path

    def carta_liberacion(self, alumno):
        doc = Document()
        self._encabezado(doc)

        doc.add_paragraph()
        titulo = doc.add_paragraph()
        titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        rt = titulo.add_run("CARTA DE LIBERACIÓN DE SERVICIO SOCIAL")
        rt.bold = True
        rt.font.size = Pt(14)
        rt.font.color.rgb = AZUL_TECNM

        doc.add_paragraph()
        fecha = datetime.now().strftime("%d de %B de %Y")
        p_fecha = doc.add_paragraph()
        p_fecha.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p_fecha.add_run(f"Iguala, Gro., a {fecha}")

        doc.add_paragraph()
        p_body = doc.add_paragraph()
        p_body.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        nombre_completo = f"{alumno['nombre']} {alumno['apellidos']}"
        texto = (
            f"Por medio de la presente, la Oficina de Servicio Social del Tecnológico Nacional de México, "
            f"hace constar que el(la) C. {nombre_completo}, con número de control "
            f"{alumno['numero_control']}, estudiante de la carrera de {alumno.get('carrera', '')}, "
            f"ha concluido satisfactoriamente su Servicio Social en el programa "
            f"\"{alumno.get('programa', '')}\", adscrito al departamento "
            f"{alumno.get('departamento', '')}, cumpliendo con las 480 horas reglamentarias "
            f"establecidas por la institución."
        )
        run = p_body.add_run(texto)
        run.font.size = Pt(12)

        doc.add_paragraph()
        doc.add_paragraph()
        doc.add_paragraph()

        p_firmas = doc.add_paragraph()
        p_firmas.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r_firma = p_firmas.add_run("_" * 35 + "          " + "_" * 35)

        p_nombres = doc.add_paragraph()
        p_nombres.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_nombres.add_run("Jefe del Departamento          Vo. Bo. Oficina Servicio Social")

        path = f"documentos/Liberacion_{alumno['numero_control']}.docx"
        doc.save(path)
        return path